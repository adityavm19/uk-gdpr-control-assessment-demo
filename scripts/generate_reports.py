from pathlib import Path
import csv
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'07-reporting'
OUT.mkdir(exist_ok=True)
DISCLAIMER='Portfolio Case Study - fictional/synthetic MTFS content only. Methodology demonstration; not a real client audit, certification audit, legal opinion or regulatory assurance engagement.'

def rows(rel):
    with open(ROOT/rel,newline='',encoding='utf-8') as f:return list(csv.DictReader(f))
controls=rows('02-assessment/article32-control-matrix.csv')
findings=rows('05-findings/findings-register.csv')
rems=rows('06-remediation/remediation-tracker.csv')
ropa=rows('02-assessment/ropa-extract.csv')

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def base_doc():
    d=Document()
    sec=d.sections[0]
    sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.65); sec.right_margin=Inches(0.65)
    styles=d.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9.5)
    styles['Title'].font.name='Aptos Display'; styles['Title'].font.size=Pt(24); styles['Title'].font.bold=True
    styles['Heading 1'].font.name='Aptos Display'; styles['Heading 1'].font.size=Pt(17); styles['Heading 1'].font.bold=True
    styles['Heading 2'].font.name='Aptos'; styles['Heading 2'].font.size=Pt(12); styles['Heading 2'].font.bold=True
    for section in d.sections:
        p=section.footer.paragraphs[0]; p.text=DISCLAIMER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.font.size=Pt(7)
    return d

def add_header(d,title):
    for section in d.sections:
        p=section.header.paragraphs[0]; p.text=title; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        for r in p.runs: r.font.size=Pt(7.5); r.font.bold=True

def add_page_number(footer_p):
    run=footer_p.add_run(' | Page ')
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text='PAGE'
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)

def add_disclaimer_box(d):
    t=d.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); c.text=DISCLAIMER; set_cell_shading(c,'F2F2F2')
    for p in c.paragraphs:
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.font.size=Pt(8); r.font.italic=True

def add_kpi_table(d, pairs):
    t=d.add_table(rows=1, cols=len(pairs)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,(label,val) in enumerate(pairs):
        c=t.cell(0,i); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(val)); r.bold=True; r.font.size=Pt(16)
        p.add_run('\n'+label)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t

def add_simple_table(d, headers, data, widths=None, font=8):
    t=d.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,'D9E2F3')
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(font)
    for row in data:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:r.font.size=Pt(font)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

def page(d): d.add_page_break()

# Executive report
ex=base_doc(); add_header(ex,'UK GDPR Security Assessment - Executive Report')
for sec in ex.sections: add_page_number(sec.footer.paragraphs[0])
p=ex.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('UK GDPR Security Control Assessment')
p=ex.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Executive report | MT-ENT-002 Meridian Trust Bank plc (fictional)'); r.bold=True; r.font.size=Pt(12)
add_disclaimer_box(ex)
ex.add_paragraph('Overall conclusion',style='Heading 1')
p=ex.add_paragraph(); p.add_run('Needs improvement. ').bold=True; p.add_run('The selected Article 32 security-control environment is broadly designed to address personal-data risk, but five operating/design weaknesses reduce assurance over privileged access, monitoring, recovery, processor oversight and vulnerability remediation. No conclusion is made on UK GDPR legal compliance as a whole.')
add_kpi_table(ex,[('control outcomes','12'),('effective','7'),('partially effective','5'),('findings','5')])
ex.add_paragraph('What matters to management',style='Heading 2')
ex.add_paragraph('Four High-rated findings affect control areas capable of increasing the likelihood or impact of unauthorised access, delayed detection, extended recovery or exploitability of customer-data systems. One Moderate finding concerns stale assurance over a material CRM processor. The underlying technical issues intentionally reuse canonical MTFS IDs already used in other assurance lenses.')
ex.add_picture(str(ROOT/'08-analytics/article32-control-coverage.png'),width=Inches(6.6))
page(ex)
ex.add_paragraph('Priority findings',style='Heading 1')
data=[]
for f in findings:
    data.append([f['finding_id'],f['title'],f['underlying_issue_id'],f['rating'],f['owner'],f['target_date']])
add_simple_table(ex,['ID','Finding','Issue','Rating','Owner','Target'],data,[0.75,2.8,0.85,0.7,1.35,0.9],7.5)
ex.add_paragraph('Risk concentration',style='Heading 2')
ex.add_picture(str(ROOT/'08-analytics/findings-by-severity.png'),width=Inches(6.5))
ex.add_paragraph('The ratings are historical assessment ratings. Planned remediation does not reduce them; closure requires evidence-backed retesting.')
page(ex)
ex.add_paragraph('Management priorities and next steps',style='Heading 1')
for f in findings:
    p=ex.add_paragraph(style=None); p.style=ex.styles['Normal']; p.add_run(f"{f['finding_id']} - {f['rating']}: ").bold=True; p.add_run(f['recommendation'])
ex.add_paragraph('Supporting-process results',style='Heading 2')
ex.add_paragraph('The breach-readiness tabletop reached a reasoned reportability decision and a hypothetical notification package 26 hours 30 minutes after awareness. The selected rights-request sample was timely under the documented workflow, including DUAA 2025 clarification timing and reasonable/proportionate-search documentation. Two restricted-transfer samples used a UK Addendum or UK IDTA plus a documented TRA/data protection test; no transfer finding was raised.')
ex.add_picture(str(ROOT/'08-analytics/remediation-status.png'),width=Inches(5.35))
ex.add_paragraph('Scope boundary',style='Heading 2')
ex.add_paragraph('This is security/control assurance supporting UK GDPR for MT-ENT-002. It excludes a full legal compliance review, PECR, EEA/DORA obligations, certification and financial-statement ITGC testing.')
ex.save(OUT/'_executive-intermediate.docx')

# Technical report
tech=base_doc(); add_header(tech,'UK GDPR Security Assessment - Technical Report')
for sec in tech.sections: add_page_number(sec.footer.paragraphs[0])
p=tech.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('UK GDPR Security Control Assessment')
p=tech.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Technical report | Report date 1 September 2026').bold=True
tech.add_paragraph(); p=tech.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Meridian Trust Bank plc (MT-ENT-002) - fictional portfolio case study').italic=True
add_disclaimer_box(tech)
tech.add_paragraph('Assessment verdict',style='Heading 1')
p=tech.add_paragraph(); p.add_run('Needs improvement. ').bold=True; p.add_run('Seven of twelve selected control outcomes were Effective and five were Partially Effective. The five partial outcomes are supported by material exceptions and are reported as findings. The assessment does not conclude on legal compliance with the UK GDPR as a whole.')
add_kpi_table(tech,[('Article 32 outcomes','12'),('Effective','7'),('Partial','5'),('Findings','5')])
tech.add_picture(str(ROOT/'08-analytics/article32-control-coverage.png'),width=Inches(6.6))
page(tech)

# page 2
tech.add_paragraph('1. Scope, boundary and criteria',style='Heading 1')
tech.add_paragraph('Engagement type: security and technical/organisational control assurance supporting UK GDPR compliance. It is not legal advice, an ICO audit, regulatory assurance or certification.')
tech.add_paragraph('Primary entity: MT-ENT-002 Meridian Trust Bank plc. Assessment period 1 April-31 July 2026; fieldwork 3-21 August 2026.')
tech.add_paragraph('In-scope assurance areas',style='Heading 2')
for x in ['Article 32 risk-appropriate technical/organisational security measures','ROPA/processing context sufficient to anchor the security assessment','Article 33-style personal-data breach readiness','Articles 15/17 rights-process workflow review','Articles 28/32 processor-security assurance','Chapter V restricted-transfer spot checks using UK safeguards']:
    tech.add_paragraph(x,style='List Bullet')
tech.add_paragraph('Explicit exclusions',style='Heading 2')
tech.add_paragraph('Full Article-by-Article legal compliance; lawful basis; privacy notices; PECR; full DPIA legal sufficiency; automated-decision legality; full retention-law analysis; penetration testing; DORA/EEA assessment; financial-audit ITGC or SAP control testing.')
tech.add_paragraph('Public criteria used',style='Heading 2')
tech.add_paragraph('UK GDPR Articles 5(1)(f), 5(2), 15, 17, 28, 32, 33 and Chapter V/Article 46 concepts, interpreted for this portfolio through current ICO guidance and the Data (Use and Access) Act 2025 changes where material.')
page(tech)

# page 3
tech.add_paragraph('2. Processing and data context',style='Heading 1')
tech.add_paragraph('Article 32 appropriateness depends on risk and processing context. Six selected processing activities anchor the assessment in customer identity, account/transaction, authentication, fraud, privacy-case and workforce-access data.')
procdata=[[r['processing_id'],r['activity'],r['data_subjects'],r['application_ids'],r['processor_ids']] for r in ropa]
add_simple_table(tech,['Processing','Activity','Data subjects','Applications','Processors'],procdata,[0.8,2.0,1.2,1.6,1.2],7.2)
tech.add_paragraph('Key canonical risk linkages',style='Heading 2')
for x in ['MT-RSK-001 Identity & Access','MT-RSK-003 Vulnerability Management','MT-RSK-004 Security Monitoring & Logging','MT-RSK-005 Backup & Recovery','MT-RSK-006 Third-Party ICT','MT-RSK-007 Data Protection & Confidentiality','MT-RSK-009 Cyber Incident Response']:
    tech.add_paragraph(x,style='List Bullet')
tech.add_paragraph('The ROPA file in this repository is a sample extract used as an assurance input. It is not presented as a legal sufficiency determination under Article 30.')
page(tech)

# page 4
tech.add_paragraph('3. Assessment methodology',style='Heading 1')
tech.add_paragraph('The methodology preserves the portfolio chain: risk -> control objective -> control -> design effectiveness -> population/sample -> evidence -> exception -> underlying issue -> GDPR finding -> remediation -> future retest.')
tech.add_paragraph('Design vs operating effectiveness',style='Heading 2')
tech.add_paragraph('Design effectiveness considers whether the defined control could meet the intended security outcome. Operating effectiveness considers whether it operated consistently during the period. Inquiry alone was not accepted as sufficient operating-effectiveness evidence where the control should leave a record.')
tech.add_paragraph('Sampling rationale',style='Heading 2')
tech.add_paragraph('Sampling is risk-directed rather than statistical. Full-population testing was used for small/high-impact populations such as 42 privileged identities, 12 critical log sources and 27 Critical/High vulnerabilities. Judgemental samples were used where full-population testing added little assurance: 20 of 99 JML events, 4 of 8 critical backup sets, 3 of 5 material processors, 12 of 74 rights cases and 2 selected restricted-transfer paths.')
tech.add_paragraph('Evidence sufficiency',style='Heading 2')
tech.add_paragraph('Evidence was evaluated for relevance, reliability, completeness, timeliness, integrity and corroboration. Twenty-two synthetic evidence artifacts are logged and hashed with SHA-256. The manifest is generated from the actual files and is validated before release.')
tech.add_paragraph('Finding rating',style='Heading 2')
tech.add_paragraph('Likelihood and impact are each scored 1-5; score is the product. Thresholds: 1-4 Low, 5-9 Moderate, 10-16 High, 17-25 Critical. Planned management action does not change the historical assessment rating.')
page(tech)

# page 5 control results
tech.add_paragraph('4. Article 32 control results',style='Heading 1')
ctl_data=[[c['control_id'],c['theme'],c['design_effectiveness'],c['operating_effectiveness'],c['sample'],c['finding_id'] or '-'] for c in controls]
add_simple_table(tech,['Control','Theme','Design','Operating','Sample','Finding'],ctl_data,[0.85,2.2,0.85,0.95,1.25,0.9],7.2)
tech.add_paragraph('Result interpretation',style='Heading 2')
tech.add_paragraph('Seven controls were Effective. Five controls were Partially Effective and produced the five findings below. No selected control was rated wholly Ineffective. This coverage is intentionally deep rather than an exhaustive checklist of every possible technical measure.')
tech.add_picture(str(ROOT/'08-analytics/article32-control-coverage.png'),width=Inches(5.6))

# next section flows naturally to the following page; avoid a blank page caused by an explicit break after a full page

tech.add_paragraph('5. Detailed findings - identity and monitoring',style='Heading 1')
for f in findings[:2]:
    tech.add_paragraph(f"{f['finding_id']} - {f['title']} ({f['rating']})",style='Heading 2')
    add_simple_table(tech,['Element','Detail'],[
        ['Underlying issue / risk',f"{f['underlying_issue_id']} | {f['risk_ids']}"],['Condition',f['condition']],['Criteria / expected state',f['criteria_expected_state']],['Cause',f['cause']],['Risk',f['risk_consequence']],['Evidence / testing',f"{f['evidence_refs']} | {f['test_workpaper_exception_refs']}"],['Recommendation',f['recommendation']],['Management action',f"{f['management_response']} Owner: {f['owner']}; target {f['target_date']}."],], [1.35,5.85],7.2)
page(tech)

# page 7 findings 3/4
tech.add_paragraph('6. Detailed findings - recovery and processor assurance',style='Heading 1')
for f in findings[2:4]:
    tech.add_paragraph(f"{f['finding_id']} - {f['title']} ({f['rating']})",style='Heading 2')
    add_simple_table(tech,['Element','Detail'],[
        ['Underlying issue / risk',f"{f['underlying_issue_id']} | {f['risk_ids']}"],['Condition',f['condition']],['Criteria / expected state',f['criteria_expected_state']],['Cause',f['cause']],['Risk',f['risk_consequence']],['Evidence / testing',f"{f['evidence_refs']} | {f['test_workpaper_exception_refs']}"],['Recommendation',f['recommendation']],['Management action',f"{f['management_response']} Owner: {f['owner']}; target {f['target_date']}."],], [1.35,5.85],7.2)
page(tech)

# page 8 finding5/remediation
tech.add_paragraph('7. Detailed finding - vulnerability ageing',style='Heading 1')
f=findings[4]
add_simple_table(tech,['Element','Detail'],[
['Finding',f"{f['finding_id']} - {f['title']} ({f['rating']}, score {f['score']})"],['Underlying issue / risk',f"{f['underlying_issue_id']} | {f['risk_ids']}"],['Condition',f['condition']],['Criteria / expected state',f['criteria_expected_state']],['Cause',f['cause']],['Risk',f['risk_consequence']],['Evidence / testing',f"{f['evidence_refs']} | {f['test_workpaper_exception_refs']}"],['Recommendation',f['recommendation']],['Management action',f"{f['management_response']} Owner: {f['owner']}; target {f['target_date']}."],], [1.35,5.85],7.2)
tech.add_paragraph('Remediation status',style='Heading 2')
remdata=[[r['remediation_id'],r['finding_id'],r['status'],r['owner'],r['target_date']] for r in rems]
add_simple_table(tech,['Action','Finding','Status','Owner','Target'],remdata,[1.0,1.0,1.2,2.1,1.0],7.2)
tech.add_picture(str(ROOT/'08-analytics/remediation-status.png'),width=Inches(6.2))
page(tech)

# page 9 supporting processes
tech.add_paragraph('8. Supporting governance exercises',style='Heading 1')
tech.add_paragraph('Breach readiness - Article 33-style tabletop',style='Heading 2')
tech.add_paragraph('Synthetic scenario: unauthorised access to 860 customer accounts with contact-data exfiltration for 143. The DPO was engaged 37 minutes after awareness; a notification decision was made the same day; the hypothetical submission package was ready after 26 hours 30 minutes. No real ICO notification occurred.')
tech.add_paragraph('Data subject rights - Articles 15 and 17 process review',style='Heading 2')
tech.add_paragraph('Population 74 Q2 cases; 12 sampled. All were timely under the documented workflow. The review explicitly considered DUAA 2025 clarification timing and the requirement to make reasonable and proportionate searches. Two terse search-rationale records were treated as an advisory documentation observation, not a material finding.')
tech.add_paragraph('Processor security - Articles 28 and 32',style='Heading 2')
tech.add_paragraph('Three of five material processors were sampled. Contractual design was adequate in all three. Northbridge operating assurance was stale, producing GDPR-FND-004.')
tech.add_paragraph('Restricted transfers - Chapter V',style='Heading 2')
tech.add_paragraph('Two US support paths were spot-checked. One used the UK Addendum and one used the UK IDTA. Both had a documented TRA/data protection test and additional security measures. The assessment does not treat EU SCCs alone as the UK transfer mechanism.')
page(tech)

# page 10 evidence + cross framework
tech.add_paragraph('9. Evidence lineage and cross-framework traceability',style='Heading 1')
tech.add_paragraph('Evidence lineage',style='Heading 2')
tech.add_paragraph('The evidence log contains 22 synthetic artifacts. Each record includes an MT-EV ID, exact relative path, SHA-256 hash, source owner/system, collection timestamp, period, canonical asset/application/service linkage and control/test references. Evidence files are immutable after logging; changed evidence receives a new ID.')
tech.add_paragraph('Canonical issue reuse',style='Heading 2')
xf=[[f['finding_id'],f['underlying_issue_id'],f['title']] for f in findings]
add_simple_table(tech,['GDPR finding','Underlying issue','Condition theme'],xf,[1.1,1.2,4.6],7.5)
tech.add_paragraph('This reuse is deliberate: the same underlying condition is not assigned a new technical issue solely because UK GDPR provides a different assurance lens. Framework-specific criteria and findings remain separate.')
tech.add_paragraph('Open finding closure rule',style='Heading 2')
tech.add_paragraph('All five findings remain Open or In Remediation. No closure or retest has been fabricated. Closure requires new evidence and a targeted retest linked to the original exception and remediation action.')
page(tech)

# page 11 sources/conclusion
tech.add_paragraph('10. Conclusion and source register',style='Heading 1')
tech.add_paragraph('Conclusion',style='Heading 2')
tech.add_paragraph('The selected control environment provides meaningful protection for personal information but requires remediation in five areas before stronger assurance is warranted. The highest priorities are privileged-access MFA consistency, vulnerability ageing, central monitoring coverage and evidenced core-banking recovery. Processor assurance should be refreshed on a risk-based schedule. Supporting breach, rights and transfer exercises were satisfactory within their stated scope.')
tech.add_paragraph('Authoritative public sources checked 1 September 2026',style='Heading 2')
source_items=[
'ICO - Security outcomes (UK GDPR security guidance).',
'ICO - Personal data breach reporting processes / 72-hour readiness.',
'ICO - A guide to subject access, updated 16 July 2026.',
'ICO - Right to erasure.',
'ICO - Records of processing and lawful basis audit framework.',
'ICO - Controller/processor responsibilities and Article 28 contract guidance.',
'ICO - International transfers, appropriate safeguards and transfer risk assessment guidance updated in 2026.',
'ICO - Data (Use and Access) Act 2025 implications for organisations, updated 19 June 2026.',
'GOV.UK - Data (Use and Access) Act 2025 data-protection/privacy changes.'
]
for x in source_items: tech.add_paragraph(x,style='List Bullet')
tech.add_paragraph('Full URLs are maintained in references/framework-sources.md and the assessment workbook. The portfolio uses public UK GDPR Article identifiers and original assessment wording; it does not reproduce proprietary standards.')
tech.add_paragraph('Final boundary statement',style='Heading 2')
tech.add_paragraph('This report demonstrates GRC/cyber-assurance methodology using fictional/synthetic information. It does not represent professional client work and does not provide a legal opinion or regulatory assurance conclusion.')
tech.save(OUT/'uk-gdpr-technical-report.docx')
print('generated reports')
